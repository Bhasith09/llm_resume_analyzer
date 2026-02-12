# pdf_generator.py

from reportlab.lib.pagesizes import letter
from reportlab.pdfgen import canvas
from reportlab.lib.utils import simpleSplit
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.enums import TA_JUSTIFY
from io import BytesIO

# For DOCX
from docx import Document
from docx.shared import Pt


# ---------------------------------------------------------
# Resume Analysis PDF (unchanged)
# ---------------------------------------------------------
def create_pdf_report(text: str) -> bytes:
    buffer = BytesIO()
    pdf = canvas.Canvas(buffer, pagesize=letter)

    width, height = letter
    y = height - 50

    pdf.setFont("Helvetica", 11)

    lines = simpleSplit(text, "Helvetica", 11, width - 80)

    for line in lines:
        if y < 50:
            pdf.showPage()
            pdf.setFont("Helvetica", 11)
            y = height - 50

        pdf.drawString(40, y, line)
        y -= 15

    pdf.save()
    buffer.seek(0)
    return buffer.read()


# ---------------------------------------------------------
# Cover Letter DOCX (Word file)
# ---------------------------------------------------------
def create_cover_letter_docx(cover_letter_text: str) -> bytes:
    doc = Document()

    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(12)

    paragraphs = cover_letter_text.split("\n")

    for para in paragraphs:
        para = para.strip()
        if para:
            p = doc.add_paragraph(para)
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(12)
        else:
            doc.add_paragraph("")

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
